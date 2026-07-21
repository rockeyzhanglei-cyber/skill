#!/usr/bin/env python3
"""
Word文档解析器
提取表格定义和红色标记的新增字段、新增表、修改字段

v3.0.0 新增：
- 新增 parse_word_document_full 函数：全量模式解析，忽略红色字体，提取所有表格
- 支持增量模式（parse_word_document）和全量模式（parse_word_document_full）

v2.5.0 新增：
- 新增 is_database_table_structure 函数：过滤非数据库表结构表格
- 只处理包含字段标识、字段名称、约束/数据类型列的表格
- 排除汇总表、指标表等非表结构定义表格

v2.4.3 修复：
- 新增各列原始值存储：required_value, format_value, data_type_category_value, data_type_value
- 用于修订记录显示实际内容（如"S2"、"AN..100"、"M"）

v2.4.0 修复：
- 约束变更方向判断：O→M不生成DDL，只有M→O才生成DDL
- 约束显示使用M/O而非"必填/可选"

v2.3.0 修复：
- 新增 is_cell_all_red 函数：检测单元格是否全部红色
- 修改 check_row_partial_red 函数：支持"某些单元格全部红色（但不是整行）"的情况
- 这种情况表示字段属性变更（如说明列的值域变更）

v2.2.0 修复：
- 新增修改字段检测：部分内容红色字体
- 区分DDL变更（约束、表示格式）和注释变更（说明、备注等）
- 新增 is_cell_partially_red、check_row_partial_red 函数

v1.21.0 修复：
- 清理从Word文档复制时带入的不可见字符（零宽空格等）
- 防止Oracle报 ORA-00911: invalid character 错误

v4.3.7 修复：
- 中文全角圆点（U+FF0E）→ 英文句点（U+002E）规范化，修复S2+N..4格式解析失败问题
- Nn 格式类型映射修正：N1/N2/Nn → VARCHAR(n)，而非 NUMBER(n)
- 固定长度的数字字符（虽然只允许数字，但类型仍是VARCHAR）

v1.17.0 修复：
- 修复红色字体检测：直接从XML元素读取color属性值
- 扩展列头映射：支持'数据元标识'、'数据元名称'等新列头
"""

import sys
import re
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

# 需要清理的不可见字符列表（从Word文档复制时可能带入）
INVISIBLE_CHARS = [
    '\u200b',  # 零宽空格 (ZERO WIDTH SPACE)
    '\u200c',  # 零宽非连接符 (ZERO WIDTH NON-JOINER)
    '\u200d',  # 零宽连接符 (ZERO WIDTH JOINER)
    '\u200e',  # 左至右标记 (LEFT-TO-RIGHT MARK)
    '\u200f',  # 右至左标记 (RIGHT-TO-LEFT MARK)
    '\u2060',  # 字连接符 (WORD JOINER)
    '\u2061',  # 函数应用 (FUNCTION APPLICATION)
    '\u2062',  # 不可见乘号 (INVISIBLE TIMES)
    '\u2063',  # 不可见分隔符 (INVISIBLE SEPARATOR)
    '\u2064',  # 不可见加号 (INVISIBLE PLUS)
    '\u206a',  # 抑止对称交换 (INHIBIT SYMMETRIC SWAPPING)
    '\u206b',  # 激活对称交换 (ACTIVATE SYMMETRIC SWAPPING)
    '\u206c',  # 抑止阿拉伯数字成形 (INHIBIT ARABIC FORM SHAPING)
    '\u206d',  # 激活阿拉伯数字成形 (ACTIVATE ARABIC FORM SHAPING)
    '\u206e',  # 国民数字形状 (NATIONAL DIGIT SHAPES)
    '\u206f',  # 欧洲数字形状 (EUROPEAN DIGIT SHAPES)
    '\ufeff',  # 零宽非断空格 (ZERO WIDTH NO-BREAK SPACE, BOM)
    '\u00ad',  # 软连字符 (SOFT HYPHEN)
]

def clean_invisible_chars(text):
    """清理文本中的不可见字符（从Word文档复制时可能带入的特殊字符）

    这些字符会导致Oracle等数据库解析报错，如 ORA-00911: invalid character

    参数:
        text: 输入文本

    返回:
        清理后的文本
    """
    if not text:
        return text
    result = text
    for char in INVISIBLE_CHARS:
        result = result.replace(char, '')
    return result

def get_color_from_xml(run):
    """从XML元素直接读取颜色值"""
    rPr = run._element.rPr
    if rPr is not None:
        color_elem = rPr.find(qn('w:color'))
        if color_elem is not None:
            val = color_elem.get(qn('w:val'))
            if val and val != 'auto':
                return val
    return None

def has_red_font(run):
    """检查run是否有红色字体（改进版：支持XML color属性）"""
    # 方法1：检查font.color.rgb（python-docx标准方式）
    if run.font.color and run.font.color.rgb:
        rgb = run.font.color.rgb
        r, g, b = rgb[0], rgb[1], rgb[2]
        # 红色范围判断 (R值高，G和B值低)
        if r > 180 and g < 80 and b < 80:
            return True

    # 方法2：直接从XML读取color属性（修复：支持w:color val="FF0000"）
    xml_color = get_color_from_xml(run)
    if xml_color:
        # 检查是否是红色（FF0000或类似红色）
        # FF0000 = 纯红, FF3333, CC0000, E60000 等都是红色系
        try:
            r = int(xml_color[0:2], 16)
            g = int(xml_color[2:4], 16)
            b = int(xml_color[4:6], 16)
            # 放宽红色判断条件：R值较高（>150），G和B值较低
            if r > 150 and g < 100 and b < 100:
                return True
            # 纯红色FF0000
            if xml_color.upper() == 'FF0000':
                return True
        except:
            pass

    return False

def check_cell_red(cell):
    """检查单元格是否有红色内容，返回红色文字片段和完整文本"""
    red_parts = []
    full_text = cell.text.strip()
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip() and has_red_font(run):
                red_parts.append(run.text.strip())
    return red_parts, full_text

def is_cell_partially_red(cell):
    """检查单元格是否部分红色（有红色内容但不是全部红色）

    返回:
        True: 部分红色（有红色内容但完整文本不全是红色）
        False: 全部红色 或 无红色内容
    """
    red_parts, full_text = check_cell_red(cell)
    if not red_parts:
        return False  # 无红色内容
    if not full_text:
        return False  # 空单元格

    # 检查是否全部红色
    red_combined = ''.join(red_parts).replace(' ', '').replace('\n', '')
    full_combined = full_text.replace(' ', '').replace('\n', '')
    if red_combined == full_combined:
        return False  # 全部红色

    return True  # 部分红色

def is_cell_all_red(cell):
    """检查单元格是否全部红色（整个单元格内容都是红色）

    返回:
        True: 全部红色（有红色内容且完整文本全是红色）
        False: 无红色内容 或 部分红色
    """
    red_parts, full_text = check_cell_red(cell)
    if not red_parts:
        return False  # 无红色内容
    if not full_text:
        return False  # 空单元格

    # 检查是否全部红色
    red_combined = ''.join(red_parts).replace(' ', '').replace('\n', '')
    full_combined = full_text.replace(' ', '').replace('\n', '')
    if red_combined == full_combined:
        return True  # 全部红色

    return False  # 部分红色

def check_row_partial_red(row, col_indices):
    """检查行是否有部分红色变更，返回变更的列信息

    v2.3.0 新增：检测"某些单元格全部红色（但不是整行）"的情况
    - 这种情况表示字段属性变更（如说明列的值域变更）
    - 约束列或表示格式列全部红色 → DDL变更
    - 说明列全部红色 → 仅注释变更

    返回:
        dict: {
            'has_constraint_change': bool,  # 约束列是否有变更
            'has_format_change': bool,      # 表示格式列是否有变更
            'has_other_change': bool,       # 其他列是否有变更（只生成注释）
            'changed_columns': list         # 变更的列名列表
            'cell_red_status': dict         # 各单元格的红色状态 {'col_key': 'all'/'partial'/'none'}
        }
    """
    result = {
        'has_constraint_change': False,
        'has_format_change': False,
        'has_other_change': False,
        'changed_columns': [],
        'cell_red_status': {}
    }

    # 需要生成DDL的列（约束和表示格式）
    ddl_columns = ['required', 'format']
    # 只生成注释的列（其他列）
    comment_columns = ['field_cn', 'comment', 'field_en', 'data_type_category', 'data_type']

    for col_key, col_idx in col_indices.items():
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            cell_text = cell.text.strip()
            if not cell_text:
                continue

            # 检查单元格的红色状态
            is_partial = is_cell_partially_red(cell)
            is_all = is_cell_all_red(cell)

            if is_partial:
                result['changed_columns'].append(col_key)
                result['cell_red_status'][col_key] = 'partial'
                if col_key in ddl_columns:
                    if col_key == 'required':
                        result['has_constraint_change'] = True
                    elif col_key == 'format':
                        result['has_format_change'] = True
                elif col_key in comment_columns:
                    result['has_other_change'] = True
            elif is_all:
                # 单元格全部红色（但不是整行红色，因为is_row_all_red已经排除了整行红色的行）
                result['changed_columns'].append(col_key)
                result['cell_red_status'][col_key] = 'all'
                if col_key in ddl_columns:
                    if col_key == 'required':
                        result['has_constraint_change'] = True
                    elif col_key == 'format':
                        result['has_format_change'] = True
                elif col_key in comment_columns:
                    result['has_other_change'] = True

    return result

def is_row_all_red(row):
    """检查整行是否红色（忽略空单元格）"""
    for cell in row.cells:
        text = cell.text.strip()
        if not text:
            continue
        red_parts, full_text = check_cell_red(cell)
        if red_parts and full_text:
            red_combined = ''.join(red_parts).replace(' ', '').replace('\n', '')
            full_combined = full_text.replace(' ', '').replace('\n', '')
            if red_combined != full_combined:
                return False
        elif not red_parts and full_text:
            return False
    return True

def is_table_all_red(table):
    """检查整个表格是否红色（新增表）"""
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if not text:
                continue
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        if not has_red_font(run):
                            return False
    return True

def get_cell_content(row, col_idx):
    """获取单元格内容，自动清理不可见字符"""
    if col_idx < len(row.cells):
        text = row.cells[col_idx].text.strip()
        return clean_invisible_chars(text)
    return ''


def parse_format_string(format_str, data_type_cat=None):
    """解析表示格式字符串，提取类型和长度

    参数:
        format_str: 表示格式字符串（如 AN..64, N1, DT19）
        data_type_cat: 数据类型类别（如 S1, S2, N, DT, D），用于区分 N1/N2 等格式的含义

    例如：
        AN..64 -> VARCHAR(64)
        N1 + 数据类型N -> NUMBER(1) 或 INTEGER
        N1 + 数据类型S -> VARCHAR(1)（数字字符型）
        DT19 -> TIMESTAMP
    """
    # 清理不可见字符
    format_str = clean_invisible_chars(format_str)
    # 中文全角圆点（U+FF0E）→ 英文句点（U+002E），Word文档中常见全角N．．4
    format_str = format_str.replace('．', '.')
    format_str = format_str.strip().upper()

    # AN系列 -> 字母数字混合（字符型）
    if format_str.startswith('AN'):
        # AN..nXm 多行（如 AN..64X3 → 行字符数 X 行数）
        match = re.search(r'AN\.{2,}(\d+)X(\d+)', format_str)
        if match:
            total = int(match.group(1)) * int(match.group(2))
            return 'VARCHAR', str(total)
        # ANn..m 范围可变长度（如 AN4..18 → 最大18）
        match = re.search(r'AN(\d+)\.{2,}(\d+)', format_str)
        if match:
            return 'VARCHAR', match.group(2)
        # AN..* → 不限制长度，映射为 TEXT（各数据库自动对应CLOB/TEXT等）
        if re.match(r'^AN\.{2,}\*$', format_str):
            return 'TEXT', ''
        # AN..n 可变长度（原逻辑）
        match = re.search(r'AN\.{2,}(\d+)', format_str)
        if match:
            return 'VARCHAR', match.group(1)
        # AN.n 单点变体
        match = re.search(r'AN\.(\d+)', format_str)
        if match:
            return 'VARCHAR', match.group(1)
        # ANn 固定长度（如 AN4）
        match = re.search(r'AN(\d+)', format_str)
        if match:
            return 'VARCHAR', match.group(1)

    # A系列 -> 纯字母字符型（规范：A = 字母字符，UTF-8编码）
    if format_str.startswith('A') and not format_str.startswith('AN'):
        # A..nXm 多行
        match = re.search(r'A\.{2,}(\d+)X(\d+)', format_str)
        if match:
            total = int(match.group(1)) * int(match.group(2))
            return 'VARCHAR', str(total)
        # A..n 可变长度
        match = re.search(r'A\.{2,}(\d+)', format_str)
        if match:
            return 'VARCHAR', match.group(1)
        # An..m 范围
        match = re.search(r'A(\d+)\.{2,}(\d+)', format_str)
        if match:
            return 'VARCHAR', match.group(2)
        # An 固定长度（如 A1, A4）
        match = re.search(r'A(\d+)(?:,(\d+))?', format_str)
        if match:
            return 'VARCHAR', match.group(1)
        # 纯 A（无长度，默认1）
        if format_str == 'A':
            return 'VARCHAR', '1'

    # T8 时间格式（规范：hh:mm:ss → 固定8字符）
    if format_str == 'T8':
        return 'VARCHAR', '8'

    # N..n 或 N..n,m -> 数值型（带精度的NUMBER类型）
    # - 数据类型 N + 表示格式 N..18,4 → NUMBER(18,4)
    # - 数据类型 N + 表示格式 N..9 → NUMBER(9)
    # - 数据类型 S3 + 表示格式 N..3 → VARCHAR(3)（代码表类型，长度是代码位数）
    if format_str.startswith('N..'):
        match = re.search(r'N\.{2,}(\d+)(?:,(\d+))?', format_str)
        if match:
            length = match.group(1)
            decimal = match.group(2) if match.group(2) else ''
            # 如果数据类型类别是 N（数值型），返回 NUMBER 类型
            if data_type_cat and data_type_cat.upper() == 'N':
                if decimal:
                    return 'NUMBER', f'{length},{decimal}'
                return 'NUMBER', length
            # 如果数据类型类别是 S1/S2/S3（字符串/代码类型），返回 VARCHAR
            if data_type_cat and data_type_cat.upper() in ['S1', 'S2', 'S3', 'S']:
                return 'VARCHAR', length
            # 如果没有明确的数据类型类别，但格式是N..开头，默认为数值型
            if decimal:
                return 'NUMBER', f'{length},{decimal}'
            return 'NUMBER', length

    # Nn -> 数值型时是数字，字符型时是数字字符
    # 关键：需要根据数据类型类别来判断
    # - 数据类型 N + 表示格式 N1 → NUMBER(1)（数值型）
    # - 数据类型 S + 表示格式 N1 → VARCHAR(1)（数字字符型）
    if format_str.startswith('N'):
        match = re.search(r'N(\d+)(?:,(\d+))?', format_str)
        if match:
            length = match.group(1)
            decimal = match.group(2) if match.group(2) else ''
            # 如果数据类型类别是 N（数值型），则是真正的数字类型
            if data_type_cat and data_type_cat.upper() == 'N':
                if decimal:
                    return 'NUMBER', f'{length},{decimal}'
                # N1, N2 等整数类型，带长度返回 NUMBER(length)
                return 'NUMBER', length
            # 否则是数字字符型（VARCHAR）
            return 'VARCHAR', length

    # DTn -> 日期时间类型（DATE，Oracle DATE可存储年月日时分秒）
    if format_str.startswith('DT'):
        return 'DATE', ''

    # Dn -> 日期类型（DATE）
    if format_str.startswith('D') and not format_str.startswith('DT'):
        return 'DATE', ''

    # S1, S2, S3 -> 字符串类型（根据约束判断）
    if format_str.startswith('S'):
        match = re.search(r'S(\d+)', format_str)
        if match:
            return 'VARCHAR', ''

    # 其他格式
    if '..' in format_str:
        # 如 AN..500, AN..2000
        parts = format_str.split('..')
        if len(parts) == 2:
            try:
                length = int(parts[1])
                return 'VARCHAR', str(length)
            except:
                pass

    return '', ''

def get_column_indices(headers):
    """根据列头获取各列索引（处理合并单元格导致的重复列头）
    v1.17.0 扩展：
    - 支持'数据元标识'、'数据元名称'等新列头
    - '数据类型'列（S1/S2/S3）只是格式类别，实际类型从'表示格式'解析
    """
    indices = {}
    seen_keys = set()  # 记录已处理的列头类型，避免重复

    for i, h in enumerate(headers):
        h_clean = h.strip().replace('\n', '')

        # 跳过空列头或已处理过的列头类型
        if not h_clean:
            continue

        # 数据元标识符/标识符/数据元标识 -> std_id（字段英文名）
        if '数据元标识符' in h_clean or '标识符' in h_clean or '数据元标识' in h_clean:
            if 'std_id' not in seen_keys:
                indices['std_id'] = i
                seen_keys.add('std_id')
                indices['field_en'] = i  # 数据元标识列包含字段英文名

        # 数据元名称/字段/数据项/字段/数据项 -> field_cn（字段中文名）
        elif '数据元名称' in h_clean or h_clean in ['字段', '数据项', '字段/数据项', '数据元名称']:
            if 'field_cn' not in seen_keys:
                indices['field_cn'] = i
                seen_keys.add('field_cn')

        # 字段名/英文名 -> field_en
        elif h_clean in ['字段名', '英文名']:
            if 'field_en' not in seen_keys:
                indices['field_en'] = i
                seen_keys.add('field_en')

        # 类型/数据类型 -> data_type_category（格式类别S1/S2/S3，不是数据库类型）
        elif h_clean in ['类型', '数据类型']:
            if 'data_type_category' not in seen_keys:
                indices['data_type_category'] = i
                seen_keys.add('data_type_category')
            # 同时也记录为data_type，以防某些文档确实有数据库类型
            if 'data_type' not in seen_keys:
                indices['data_type'] = i
                seen_keys.add('data_type')

        # 长度/数据长度 -> length
        elif h_clean in ['长度', '数据长度']:
            if 'length' not in seen_keys:
                indices['length'] = i
                seen_keys.add('length')

        # 表示格式 -> format（用于解析实际数据库类型和长度）
        elif '表示格式' in h_clean or h_clean == '表示格式':
            if 'format' not in seen_keys:
                indices['format'] = i
                seen_keys.add('format')

        # 填报要求/必填/约束 -> required
        elif '填报' in h_clean or '必填' in h_clean or h_clean == '约束':
            if 'required' not in seen_keys:
                indices['required'] = i
                seen_keys.add('required')

        # 说明/备注/值域 -> comment
        elif h_clean in ['说明', '备注', '值域']:
            if 'comment' not in seen_keys:
                indices['comment'] = i
                seen_keys.add('comment')

    # 如果没有找到field_en，使用std_id
    if 'field_en' not in indices and 'std_id' in indices:
        indices['field_en'] = indices['std_id']

    # 如果没有找到field_cn，查找std_id后面的一列
    if 'field_cn' not in indices:
        if 'std_id' in indices:
            # 查找std_id后面且不是已知列头类型的位置
            for i, h in enumerate(headers):
                h_clean = h.strip().replace('\n', '')
                if i > indices.get('std_id', -1) and h_clean not in ['', '字段名', '类型', '长度', '填报要求', '说明', '约束', '数据类型', '表示格式']:
                    indices['field_cn'] = i
                    break
        if 'field_cn' not in indices:
            indices['field_cn'] = 1  # 默认第二列

    return indices

def extract_table_names_ordered(doc):
    """按文档顺序提取表名，关联表格位置"""
    table_names = []
    body_elements = doc.element.body
    current_para_idx = 0
    current_table_idx = 0

    for element in body_elements:
        if element.tag.endswith('p'):  # 段落
            para = doc.paragraphs[current_para_idx]
            # 清理段落文本中的不可见字符
            text = clean_invisible_chars(para.text.strip())

            # 只处理标题样式的段落（排除普通文本格式的表说明）
            # Heading样式才是真正的表名，Normal样式是说明文字
            style_name = para.style.name if para.style else 'Normal'
            is_heading = style_name.startswith('Heading') or '标题' in style_name
            
            if is_heading:
                # 检查是否包含表名（支持两种格式）
                # 格式1：中文名 EN_TABLE_NAME（表名在末尾，空格分隔）
                # 格式2：中文名（EN_TABLE_NAME）（表名在中文括号内）
                match = re.search(r'([A-Za-z][A-Za-z0-9_]{3,})\s*$', text)
                if not match:
                    # 尝试匹配中文括号格式
                    match = re.search(r'[（\(]([A-Za-z][A-Za-z0-9_]{3,})[）\)]', text)
            
                if match:
                    table_en = clean_invisible_chars(match.group(1))
                    # 提取中文部分（去掉括号内的表名和末尾的英文表名）
                    text_for_cn = re.sub(r'[（\(][A-Za-z][A-Za-z0-9_]+[）\)]', '', text).strip()
                    text_for_cn = re.sub(r'[A-Za-z][A-Za-z0-9_]+\s*$', '', text_for_cn).strip()
                    cn_match = re.search(r'(.+?表|.+?信息|.+?记录|.+?清单|.+?目录)', text_for_cn)
                    if cn_match:
                        table_cn = clean_invisible_chars(cn_match.group(1).strip())
                    else:
                        # 如果没有匹配到关键字，直接使用去掉表名后的文本
                        table_cn = text_for_cn

                    if table_cn and table_en:
                        # 直接检查段落中的run是否红色
                        is_para_red = False
                        for run in para.runs:
                            if run.text.strip() and has_red_font(run):
                                is_para_red = True
                                break

                        table_names.append({
                            'table_en': table_en,
                            'table_cn': table_cn,
                            'para_idx': current_para_idx,
                            'pending': True,
                            'is_new_table': is_para_red  # 红色段落标题表示新增表
                        })

            current_para_idx += 1

        elif element.tag.endswith('tbl'):  # 表格
            # 关联最近pending的表名
            for tn in reversed(table_names):
                if tn.get('pending'):
                    tn['table_idx'] = current_table_idx
                    tn['pending'] = False

                    # 同时检查表格本身是否整表红色
                    table = doc.tables[current_table_idx]
                    if is_table_all_red(table):
                        tn['is_new_table'] = True

                    break

            current_table_idx += 1

    return table_names

def extract_categories(doc):
    """提取数据集区域分类目录（二级标题），检测是否新增（红色字体）

    只有二级标题下面有数据集表格时，才算是真正的数据集分类。
    数据集文档结构：二级标题(分类) → 三级标题(数据集名称) → 数据集表格

    返回:
        list: [{'category_name': 'xxx', 'is_new': True/False, 'para_idx': n}]
    """
    categories = []
    body_elements = doc.element.body

    # 第一步：建立段落索引到element索引的映射，并找到所有表格的位置
    # para_to_element[para_idx] = element_idx
    para_to_element = {}
    table_positions = []  # [{'element_idx': n, 'is_dataset_table': True/False}]
    para_idx = 0
    tbl_idx = 0
    element_idx = 0

    for element in body_elements:
        if element.tag.endswith('p'):  # 段落
            para_to_element[para_idx] = element_idx
            para_idx += 1
        elif element.tag.endswith('tbl'):  # 表格
            # 检查表格是否是数据集表格
            table = doc.tables[tbl_idx]
            if len(table.rows) > 0:
                header_row = table.rows[0]
                col_indices = {}
                for i, cell in enumerate(header_row.cells):
                    cell_text = clean_invisible_chars(cell.text.strip())
                    # 检测列头
                    if '数据元标识' in cell_text or '数据元标识符' in cell_text or '字段名' in cell_text or '英文名' in cell_text:
                        col_indices['std_id'] = i
                    elif '数据元名称' in cell_text or '字段' in cell_text or '数据项' in cell_text:
                        col_indices['field_cn'] = i
                    elif '约束' in cell_text or '填报要求' in cell_text:
                        col_indices['required'] = i
                    elif '表示格式' in cell_text:
                        col_indices['format'] = i
                    elif '数据类型' in cell_text:
                        col_indices['data_type'] = i

                is_dataset_table = is_database_table_structure(col_indices)
                table_positions.append({'element_idx': element_idx, 'is_dataset_table': is_dataset_table})
            tbl_idx += 1
        element_idx += 1

    # 第二步：遍历paragraphs，找到二级标题
    potential_categories = []

    for p_idx in range(len(doc.paragraphs)):
        para = doc.paragraphs[p_idx]
        text = clean_invisible_chars(para.text.strip())

        # 跳过空段落和表格标题（包含英文表名）
        if not text or re.search(r'[A-Za-z][A-Za-z0-9_]{3,}\s*$', text):
            continue

        # 检查段落样式是否是二级标题（Heading 2或标题 2）
        style_name = para.style.name if para.style else ''
        is_heading2 = style_name == 'Heading 2' or style_name == '标题 2' or 'Heading 2' in style_name or '标题 2' in style_name

        if not is_heading2:
            continue

        # 检查是否是分类标题（不含英文表名）
        # 分类标题：纯中文，不含英文表名（如"不良事件"、"临床路径"）
        # 数据集标题：包含英文表名（如"不良事件_严重不良事件报告 AE_SERIOUS"）
        has_english_table_name = re.search(r'[A-Za-z][A-Za-z0-9_]{3,}', text)

        if not has_english_table_name and len(text) < 50:
            # 检查段落是否红色
            is_para_red = False
            for run in para.runs:
                if run.text.strip() and has_red_font(run):
                    is_para_red = True
                    break

            # 获取该段落在body_elements中的位置
            element_idx = para_to_element.get(p_idx, -1)

            potential_categories.append({
                'category_name': text,
                'is_new': is_para_red,
                'para_idx': p_idx,
                'element_idx': element_idx
            })

    # 第三步：对于每个二级标题，检查它后面是否有数据集表格（直到遇到下一个二级标题）
    for i, cat in enumerate(potential_categories):
        cat_element_idx = cat['element_idx']
        # 找到下一个二级标题的位置（如果没有，则到文档末尾）
        next_heading2_element_idx = None
        for j in range(i + 1, len(potential_categories)):
            if potential_categories[j]['element_idx'] > cat_element_idx:
                next_heading2_element_idx = potential_categories[j]['element_idx']
                break

        # 检查在当前二级标题和下一个二级标题之间是否有数据集表格
        has_dataset_table = False
        for tbl in table_positions:
            tbl_element_idx = tbl['element_idx']
            if tbl_element_idx > cat_element_idx:
                if next_heading2_element_idx is None or tbl_element_idx < next_heading2_element_idx:
                    if tbl['is_dataset_table']:
                        has_dataset_table = True
                        break

        # 只有后面有数据集表格的二级标题才是真正的数据集分类
        if has_dataset_table:
            categories.append(cat)

    return categories

def is_database_table_structure(col_indices):
    """判断表格是否是数据库表结构定义

    数据库表结构表格必须同时有以下列：
    - 字段标识列（std_id 或 field_en）
    - 字段名称列（field_cn）
    - 约束或数据类型列（required、format 或 data_type）

    参数:
        col_indices: 列索引字典

    返回:
        True: 是数据库表结构
        False: 非数据库表结构（如汇总表、指标表等）
    """
    # 必须有字段标识列
    has_field_id = 'std_id' in col_indices or 'field_en' in col_indices

    # 必须有字段名称列
    has_field_name = 'field_cn' in col_indices

    # 必须有约束或数据类型列
    has_type_or_constraint = (
        'required' in col_indices or
        'format' in col_indices or
        'data_type' in col_indices or
        'data_type_category' in col_indices
    )

    return has_field_id and has_field_name and has_type_or_constraint

def parse_word_document(doc_path):
    """解析Word文档，提取所有有变更的表格和新增表

    返回结果包含：
    - new_tables: 新增表列表
    - all_changes: 新增字段的表列表
    - modified_fields: 修改字段的表列表（区分DDL变更和注释变更）
    - new_categories: 新增分类列表（红色字体的二级标题）
    - doc_cover: 文档封面信息（标准名称）
    """

    doc = Document(doc_path)
    tables = doc.tables

    # 提取分类目录
    categories = extract_categories(doc)
    new_categories = [c for c in categories if c['is_new']]

    # 提取文档封面信息（第一个段落通常是封面）
    doc_cover = ''
    if doc.paragraphs:
        first_para_text = clean_invisible_chars(doc.paragraphs[0].text.strip())
        if first_para_text and len(first_para_text) > 10:
            doc_cover = first_para_text[:100]  # 截取前100字符作为封面信息

    # 提取表名映射
    table_name_list = extract_table_names_ordered(doc)

    # 构建分类索引（用于关联数据集）
    category_idx = 0
    current_category = None
    table_idx_to_category = {}

    # 根据段落顺序，找到每个表格对应的分类（最近的分类标题）
    # 先按段落索引排序分类
    sorted_categories = sorted(categories, key=lambda x: x['para_idx'])

    # 遍历所有表格，找到每个表格之前最近的分类
    for tn in table_name_list:
        if 'table_idx' in tn and 'para_idx' in tn:
            table_para_idx = tn['para_idx']
            # 找到表格之前最近的分类
            best_category = ''
            for cat in sorted_categories:
                if cat['para_idx'] < table_para_idx:
                    best_category = cat['category_name']
                else:
                    break  # 分类在表格后面，停止搜索
            table_idx_to_category[tn['table_idx']] = best_category

    table_idx_to_name = {}
    new_tables = []  # 新增表列表

    for tn in table_name_list:
        if 'table_idx' in tn:
            category_name = table_idx_to_category.get(tn['table_idx'], '')
            table_idx_to_name[tn['table_idx']] = {
                'cn': tn['table_cn'],
                'en': tn['table_en'],
                'is_new': tn.get('is_new_table', False),
                'category_name': category_name
            }
            # 记录新增表
            if tn.get('is_new_table'):
                new_tables.append({
                    'table_cn': tn['table_cn'],
                    'table_en': tn['table_en'],
                    'table_idx': tn['table_idx'],
                    'category_name': category_name
                })

    # 分析所有变更
    all_changes = []  # 新增字段的表
    modified_fields = []  # 修改字段的表
    detected_headers = set()

    for table_idx, table in enumerate(tables):
        if len(table.rows) < 2:
            continue

        headers = [cell.text.strip().replace('\n', '') for cell in table.rows[0].cells]
        detected_headers.update(headers)
        col_indices = get_column_indices(headers)

        # v2.5.0: 过滤非数据库表结构表格
        # 只处理包含字段标识、字段名称、约束/数据类型列的表格
        if not is_database_table_structure(col_indices):
            continue

        # 跳过新增表（新增表生成CREATE TABLE，不是ALTER TABLE）
        if table_idx in table_idx_to_name and table_idx_to_name[table_idx]['is_new']:
            continue

        # 检查红色行（新增字段 - 整行红色）
        red_rows = []
        # 检查部分红色行（修改字段 - 部分内容红色）
        partial_red_rows = []

        for row_idx, row in enumerate(table.rows[1:], start=1):
            if is_row_all_red(row):
                red_rows.append(row_idx)
            else:
                # 检查是否有部分红色变更
                partial_result = check_row_partial_red(row, col_indices)
                if partial_result['changed_columns']:
                    partial_red_rows.append((row_idx, partial_result))

        if not red_rows and not partial_red_rows:
            continue

        # 提取新增字段（整行红色）
        new_fields = []
        for row_idx in red_rows:
            row = table.rows[row_idx]

            field_cn = get_cell_content(row, col_indices.get('field_cn', col_indices.get('std_id', 0) + 1))
            field_en = get_cell_content(row, col_indices.get('field_en', col_indices.get('std_id', 0)))
            data_type_col = get_cell_content(row, col_indices.get('data_type', 3))  # 可能是S1/S2/S3或实际类型
            format_col = get_cell_content(row, col_indices.get('format', 4))  # 表示格式列（优先使用）
            length_col = get_cell_content(row, col_indices.get('length', 3))
            required = get_cell_content(row, col_indices.get('required', 2))  # 约束列
            comment = get_cell_content(row, col_indices.get('comment', 5))

            # 清理字段名中的不可见字符和括号
            field_en_clean = clean_invisible_chars(re.sub(r'[（）\(\)]', '', field_en).strip())

            # 解析数据类型和长度
            # 优先从"表示格式"列解析（AN..64/N1/DT19），因为"数据类型"列可能只是格式类别S1/S2/S3
            # 注意：需要传入数据类型类别，以便正确区分 N1 是数值型还是数字字符型
            data_type = ''
            length = ''

            if format_col:
                parsed_type, parsed_length = parse_format_string(format_col, data_type_col)
                if parsed_type:
                    data_type = parsed_type
                    length = parsed_length

            # 如果表示格式没解析出类型，再检查数据类型列是否是实际数据库类型
            if not data_type and data_type_col:
                # 判断是否是实际数据库类型（VARCHAR/NUMBER/DATE等）而非格式类别（S1/S2/S3）
                if data_type_col.upper() not in ['S1', 'S2', 'S3', 'S', 'S4', 'N', 'DT', 'D']:
                    data_type = data_type_col.strip()
                    if length_col:
                        length = length_col.strip()

            # 如果还是没有类型，尝试从数据类型类别推断
            if not data_type and data_type_col:
                cat = data_type_col.upper()
                if cat in ['S1', 'S2']:  # 字符串类型
                    data_type = 'VARCHAR'
                elif cat == 'S3':  # 代码表类型
                    data_type = 'VARCHAR'
                elif cat == 'DT':  # 日期时间类型
                    data_type = 'DATE'
                elif cat == 'D':  # 日期类型
                    data_type = 'DATE'
                elif cat == 'N':  # 数值类型
                    data_type = 'NUMBER'

            if field_en_clean and data_type:
                # 约束判断：根据原始值确定显示文本
                # v4.3.4: 条件必填显示"条件必填"，空白显示"应填"，M显示"必填"
                required_lower = required.lower() if required else ''
                required_cn = '应填'  # 默认值
                
                if '条件必填' in required or '条件' in required_lower:
                    required_cn = '条件必填'
                elif required.strip() == '' or required_lower in ['', 'o', '可选']:
                    required_cn = '应填'
                elif 'm' in required_lower or '必填' in required:
                    required_cn = '必填'
                
                # 脚本中新增字段统一使用null（非必填），避免已有数据插入失败
                constraint = ''

                new_fields.append({
                    'field_cn': field_cn,
                    'field_en': field_en_clean,
                    'data_type': data_type,
                    'length': length,
                    'constraint': constraint,
                    'required_cn': required_cn,
                    'comment': comment,
                    # v3.0.1: 新增原始值保存，用于DML生成
                    'required_value': required,
                    'format_value': format_col,
                    'data_type_value': data_type_col,
                    'data_type_category_value': data_type_col
                })

        # 提取修改字段（部分红色）
        modified_fields_for_table = []
        for row_idx, partial_result in partial_red_rows:
            row = table.rows[row_idx]

            field_cn = get_cell_content(row, col_indices.get('field_cn', col_indices.get('std_id', 0) + 1))
            field_en = get_cell_content(row, col_indices.get('field_en', col_indices.get('std_id', 0)))
            format_col = get_cell_content(row, col_indices.get('format', 4))
            length_col = get_cell_content(row, col_indices.get('length', 3))
            data_type_col = get_cell_content(row, col_indices.get('data_type', 3))
            data_type_category_col = get_cell_content(row, col_indices.get('data_type_category', 3))
            required = get_cell_content(row, col_indices.get('required', 2))
            comment = get_cell_content(row, col_indices.get('comment', 5))

            # 清理字段名中的不可见字符和括号
            field_en_clean = clean_invisible_chars(re.sub(r'[（）\(\)]', '', field_en).strip())

            # 解析当前数据类型和长度
            data_type = ''
            length = ''

            if format_col:
                parsed_type, parsed_length = parse_format_string(format_col, data_type_col)
                if parsed_type:
                    data_type = parsed_type
                    length = parsed_length

            if not data_type and data_type_col:
                if data_type_col.upper() not in ['S1', 'S2', 'S3', 'S', 'S4', 'N', 'DT', 'D']:
                    data_type = data_type_col.strip()
                    if length_col:
                        length = length_col.strip()

            if not data_type and data_type_col:
                cat = data_type_col.upper()
                if cat in ['S1', 'S2']:
                    data_type = 'VARCHAR'
                elif cat == 'S3':
                    data_type = 'VARCHAR'
                elif cat == 'DT':
                    data_type = 'TIMESTAMP'
                elif cat == 'D':
                    data_type = 'DATE'
                elif cat == 'N':
                    data_type = 'NUMBER'

            # 约束判断：根据原始值确定显示文本
            # v4.3.4: 条件必填显示"条件必填"，空白显示"应填"，M显示"必填"
            required_lower = required.lower() if required else ''
            required_cn = '应填'  # 默认值
            
            if '条件必填' in required or '条件' in required_lower:
                required_cn = '条件必填'
                is_required = False  # 条件必填按非必填处理
            elif required.strip() == '' or required_lower in ['', 'o', '可选']:
                required_cn = '应填'
                is_required = False
            elif 'm' in required_lower or '必填' in required:
                required_cn = '必填'
                is_required = True
            else:
                is_required = False
            
            # 脚本中修改字段也统一使用null
            constraint = ''

            # v2.4.0: 判断约束变更方向
            # 当约束列变红时，当前显示的值是"新值"
            # 如果新值是M（必填），则变更方向是 O→M，不生成DDL
            # 如果新值是O（可选），则变更方向是 M→O，需要生成DDL
            constraint_change_direction = ''
            if partial_result['has_constraint_change']:
                if is_required:
                    constraint_change_direction = 'O_to_M'  # 从可选改为必填，不生成DDL
                else:
                    constraint_change_direction = 'M_to_O'  # 从必填改为可选，生成DDL

            # 构建修改字段信息
            # v2.4.0: 约束改为M（O→M）时，不生成DDL
            # 只有 M→O 或表示格式变更才生成DDL
            has_real_ddl_change = False
            if partial_result['has_format_change']:
                has_real_ddl_change = True
            if partial_result['has_constraint_change'] and constraint_change_direction == 'M_to_O':
                has_real_ddl_change = True

            mod_field = {
                'field_cn': field_cn,
                'field_en': field_en_clean,
                'data_type': data_type,
                'length': length,
                'constraint': constraint,
                'required_cn': required_cn,  # v4.3.4: 使用新的约束显示逻辑
                'comment': comment,
                'changed_columns': partial_result['changed_columns'],
                'has_ddl_change': has_real_ddl_change,  # v2.4.0: 使用新的DDL判断逻辑
                'has_constraint_change': partial_result['has_constraint_change'],
                'has_format_change': partial_result['has_format_change'],
                'has_other_change': partial_result['has_other_change'],
                'constraint_change_direction': constraint_change_direction,  # v2.4.0: 新增
                # v2.4.2: 新增各列原始值存储，用于修订记录显示
                'required_value': required,  # 约束列原始值 (M/O)
                'format_value': format_col,  # 表示格式列原始值 (AN..200等)
                'data_type_category_value': data_type_category_col,  # 数据类型类别列原始值 (S2等)
                'data_type_value': data_type_col  # 数据类型列原始值
            }

            modified_fields_for_table.append(mod_field)

        # 收集变更
        if new_fields and table_idx in table_idx_to_name:
            table_info = table_idx_to_name[table_idx]
            category_name = table_idx_to_category.get(table_idx, '')
            all_changes.append({
                'table_idx': table_idx,
                'table_cn': table_info['cn'],
                'table_en': table_info['en'],
                'category_name': category_name,
                'new_fields': new_fields
            })

        if modified_fields_for_table and table_idx in table_idx_to_name:
            table_info = table_idx_to_name[table_idx]
            category_name = table_idx_to_category.get(table_idx, '')
            modified_fields.append({
                'table_idx': table_idx,
                'table_cn': table_info['cn'],
                'table_en': table_info['en'],
                'category_name': category_name,
                'modified_fields': modified_fields_for_table
            })

    # 分析新增表的字段（用于生成CREATE TABLE）
    new_table_details = []
    for nt in new_tables:
        table_idx = nt['table_idx']
        if table_idx < len(tables):
            table = tables[table_idx]
            if len(table.rows) < 2:
                continue

            headers = [cell.text.strip().replace('\n', '') for cell in table.rows[0].cells]
            col_indices = get_column_indices(headers)

            # 提取所有字段，同时检测主键
            all_fields = []
            primary_key_fields = []
            for row_idx, row in enumerate(table.rows[1:], start=1):
                field_cn = get_cell_content(row, col_indices.get('field_cn', col_indices.get('std_id', 0) + 1))
                field_en = get_cell_content(row, col_indices.get('field_en', col_indices.get('std_id', 0)))
                data_type_col = get_cell_content(row, col_indices.get('data_type', 3))  # 可能是S1/S2/S3或实际类型
                format_col = get_cell_content(row, col_indices.get('format', 4))  # 表示格式列（优先使用）
                length_col = get_cell_content(row, col_indices.get('length', 3))
                required = get_cell_content(row, col_indices.get('required', 2))
                comment = get_cell_content(row, col_indices.get('comment', 5))

                # 清理字段名中的不可见字符和括号
                field_en_clean = clean_invisible_chars(re.sub(r'[（）\(\)]', '', field_en).strip())

                # 解析数据类型和长度
                # 优先从"表示格式"列解析（AN..64/N1/DT19），因为"数据类型"列可能只是格式类别S1/S2/S3
                data_type = ''
                length = ''

                if format_col:
                    parsed_type, parsed_length = parse_format_string(format_col, data_type_col)
                    if parsed_type:
                        data_type = parsed_type
                        length = parsed_length

                # 如果表示格式没解析出类型，再检查数据类型列是否是实际数据库类型
                if not data_type and data_type_col:
                    # 判断是否是实际数据库类型（VARCHAR/NUMBER/DATE等）而非格式类别（S1/S2/S3）
                    if data_type_col.upper() not in ['S1', 'S2', 'S3', 'S', 'S4', 'N', 'DT', 'D']:
                        data_type = data_type_col.strip()
                        if length_col:
                            length = length_col.strip()

                # 如果还是没有类型，尝试从数据类型类别推断
                if not data_type and data_type_col:
                    cat = data_type_col.upper()
                    if cat in ['S1', 'S2', 'S3']:  # 都是字符型
                        data_type = 'VARCHAR'
                    elif cat == 'DT':  # 日期时间类型
                        data_type = 'TIMESTAMP'
                    elif cat == 'D':  # 日期类型
                        data_type = 'DATE'
                    elif cat == 'N':  # 数字类型
                        data_type = 'NUMBER'
                    elif cat == 'L':  # 布尔型
                        data_type = 'VARCHAR'  # 或 NUMBER(1)

                if field_en_clean and data_type:
                    is_required = 'M' in required.upper() or '必填' in required
                    constraint = 'not null' if is_required else ''

                    # 检测主键（说明列包含"复合主键"或"联合主键"）
                    is_pk = False
                    if comment and ('复合主键' in comment or '联合主键' in comment or '主键' in comment):
                        is_pk = True
                        primary_key_fields.append(field_en_clean)

                    all_fields.append({
                        'field_cn': field_cn,
                        'field_en': field_en_clean,
                        'data_type': data_type,
                        'length': length,
                        'constraint': constraint,
                        'required_cn': '必填' if is_required else '应填',
                        'comment': comment,
                        'is_pk': is_pk,
                        # v3.0.1: 新增原始值保存，用于DML生成
                        'required_value': required,
                        'format_value': format_col,
                        'data_type_value': data_type_col,
                        'data_type_category_value': data_type_col,
                        'representation_format': format_col  # DML使用原始值
                    })

            new_table_details.append({
                'table_cn': nt['table_cn'],
                'table_en': nt['table_en'],
                'category_name': nt.get('category_name', ''),
                'fields': all_fields,
                'primary_keys': primary_key_fields
            })

    return {
        'total_tables': len(tables),
        'detected_headers': sorted(list(detected_headers)),
        'new_tables': new_table_details,
        'changed_tables': len(all_changes),
        'all_changes': all_changes,
        'modified_tables': len(modified_fields),
        'modified_fields': modified_fields,
        'new_categories': new_categories,
        'doc_cover': doc_cover,
        'table_idx_to_category': table_idx_to_category
    }

def parse_word_document_full(doc_path):
    """全量模式解析Word文档，忽略红色字体，提取所有表格作为新增表

    v3.0.0 新增：用于全量DDL/DML生成场景

    返回结果包含：
    - new_tables: 所有表格作为新增表列表
    - new_categories: 所有分类（忽略红色标记）
    - doc_cover: 文档封面信息（标准名称）
    - 注意：all_changes 和 modified_fields 为空（全量模式不区分）
    """
    doc = Document(doc_path)
    tables = doc.tables

    # 提取分类目录（全量模式：所有分类都视为已存在，不标记为新增）
    categories = extract_categories(doc)
    # 全量模式下，分类标记为非新增（is_new=False）
    all_categories = [{'category_name': c['category_name'], 'is_new': False, 'para_idx': c['para_idx']} for c in categories]

    # 提取文档封面信息
    doc_cover = ''
    if doc.paragraphs:
        first_para_text = clean_invisible_chars(doc.paragraphs[0].text.strip())
        if first_para_text and len(first_para_text) > 10:
            doc_cover = first_para_text[:100]

    # 提取表名映射（全量模式：忽略红色标记）
    table_name_list = extract_table_names_ordered(doc)

    # 构建分类索引
    sorted_categories = sorted(categories, key=lambda x: x['para_idx'])
    table_idx_to_category = {}

    for tn in table_name_list:
        if 'table_idx' in tn and 'para_idx' in tn:
            table_para_idx = tn['para_idx']
            best_category = ''
            for cat in sorted_categories:
                if cat['para_idx'] < table_para_idx:
                    best_category = cat['category_name']
                else:
                    break
            table_idx_to_category[tn['table_idx']] = best_category

    # 全量模式：所有表格都视为新增表
    new_table_details = []
    table_idx_to_name = {}

    for tn in table_name_list:
        if 'table_idx' in tn:
            category_name = table_idx_to_category.get(tn['table_idx'], '')
            table_idx_to_name[tn['table_idx']] = {
                'cn': tn['table_cn'],
                'en': tn['table_en'],
                'category_name': category_name
            }

    # 解析所有表格的字段
    for table_idx, table in enumerate(tables):
        if len(table.rows) < 2:
            continue

        headers = [cell.text.strip().replace('\n', '') for cell in table.rows[0].cells]
        col_indices = get_column_indices(headers)

        # 过滤非数据库表结构表格
        if not is_database_table_structure(col_indices):
            continue

        # 查找表名信息
        if table_idx not in table_idx_to_name:
            continue

        table_info = table_idx_to_name[table_idx]
        category_name = table_idx_to_category.get(table_idx, '')

        # 提取所有字段
        all_fields = []
        primary_key_fields = []

        for row_idx, row in enumerate(table.rows[1:], start=1):
            field_cn = get_cell_content(row, col_indices.get('field_cn', col_indices.get('std_id', 0) + 1))
            field_en = get_cell_content(row, col_indices.get('field_en', col_indices.get('std_id', 0)))
            data_type_col = get_cell_content(row, col_indices.get('data_type', 3))
            format_col = get_cell_content(row, col_indices.get('format', 4))
            length_col = get_cell_content(row, col_indices.get('length', 3))
            required = get_cell_content(row, col_indices.get('required', 2))
            comment = get_cell_content(row, col_indices.get('comment', 5))

            field_en_clean = clean_invisible_chars(re.sub(r'[（）\(\)]', '', field_en).strip())

            # 解析数据类型和长度
            data_type = ''
            length = ''

            if format_col:
                parsed_type, parsed_length = parse_format_string(format_col, data_type_col)
                if parsed_type:
                    data_type = parsed_type
                    length = parsed_length

            if not data_type and data_type_col:
                if data_type_col.upper() not in ['S1', 'S2', 'S3', 'S', 'S4', 'N', 'DT', 'D']:
                    data_type = data_type_col.strip()
                    if length_col:
                        length = length_col.strip()

            if not data_type and data_type_col:
                cat = data_type_col.upper()
                if cat in ['S1', 'S2', 'S3']:
                    data_type = 'VARCHAR'
                elif cat == 'DT':
                    data_type = 'DATE'
                elif cat == 'D':
                    data_type = 'DATE'
                elif cat == 'N':
                    data_type = 'NUMBER'
                elif cat == 'L':
                    data_type = 'VARCHAR'

            if field_en_clean and data_type:
                is_required = 'M' in required.upper() or '必填' in required
                constraint = 'not null' if is_required else ''

                is_pk = False
                if comment and ('复合主键' in comment or '联合主键' in comment or '主键' in comment):
                    is_pk = True
                    primary_key_fields.append(field_en_clean)

                all_fields.append({
                    'field_cn': field_cn,
                    'field_en': field_en_clean,
                    'data_type': data_type,
                    'length': length,
                    'constraint': constraint,
                    'required_cn': '必填' if is_required else '应填',
                    'comment': comment,
                    'is_pk': is_pk,
                    'required_value': required,
                    'format_value': format_col,
                    'data_type_value': data_type_col,
                    'data_type_category_value': data_type_col,
                    'representation_format': format_col
                })

        if all_fields:
            new_table_details.append({
                'table_cn': table_info['cn'],
                'table_en': table_info['en'],
                'category_name': category_name,
                'fields': all_fields,
                'primary_keys': primary_key_fields
            })

    return {
        'total_tables': len(tables),
        'new_tables': new_table_details,
        'changed_tables': 0,  # 全量模式下无"新增字段"概念
        'all_changes': [],    # 全量模式下无"新增字段"概念
        'modified_tables': 0, # 全量模式下无"修改字段"概念
        'modified_fields': [],# 全量模式下无"修改字段"概念
        'new_categories': [], # 全量模式下分类不标记为新增
        'all_categories': all_categories,  # 所有分类（用于全量DML）
        'doc_cover': doc_cover,
        'table_idx_to_category': table_idx_to_category,
        'mode': 'full'  # 标识为全量模式
    }


def main():
    """主函数，用于命令行调用"""
    if len(sys.argv) < 2:
        print("用法: python word_parser.py <word文档路径>")
        sys.exit(1)

    doc_path = sys.argv[1]
    result = parse_word_document(doc_path)

    print(f"文档共有 {result['total_tables']} 个表格")
    print(f"检测到列头: {result['detected_headers'][:10]}")
    print(f"新增分类: {len(result['new_categories'])} 个")
    print(f"新增表: {len(result['new_tables'])} 个")
    print(f"新增字段的表: {result['changed_tables']} 个")
    print(f"修改字段的表: {result['modified_tables']} 个")

    # 输出新增分类
    if result['new_categories']:
        print("\n=== 新增分类 ===")
        for cat in result['new_categories'][:10]:
            print(f"  + {cat['category_name']}")

    # 输出新增表
    if result['new_tables']:
        print("\n=== 新增表 ===")
        for nt in result['new_tables']:
            print(f"新增表：{nt['table_cn']}[{nt['table_en']}] ({len(nt['fields'])} 个字段)")
            if nt.get('category_name'):
                print(f"  分类: {nt['category_name']}")
            for f in nt['fields'][:5]:
                print(f"  - {f['field_cn']} [{f['field_en']}, {f['data_type']}, {f['required_cn']}]")

    # 输出新增字段
    print("\n=== 新增字段 ===")
    for change in result['all_changes'][:5]:
        print(f"\n{change['table_cn']}[{change['table_en']}]:")
        if change.get('category_name'):
            print(f"  分类: {change['category_name']}")
        for f in change['new_fields'][:3]:
            print(f"  + {f['field_cn']} [{f['field_en']}, {f['data_type']}, {f['required_cn']}]")

    # 输出修改字段
    if result['modified_fields']:
        print("\n=== 修改字段 ===")
        for mod in result['modified_fields'][:5]:
            print(f"\n{mod['table_cn']}[{mod['table_en']}]:")
            if mod.get('category_name'):
                print(f"  分类: {mod['category_name']}")
            for f in mod['modified_fields'][:3]:
                ddl_flag = "★DDL" if f['has_ddl_change'] else "仅注释"
                changed_cols = ', '.join(f['changed_columns'])
                print(f"  ~ {f['field_cn']} [{f['field_en']}] - 修改属性: {changed_cols} ({ddl_flag})")

if __name__ == '__main__':
    main()
