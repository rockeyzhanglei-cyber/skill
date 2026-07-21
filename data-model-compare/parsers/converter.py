#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档格式转换器 - 简单模式
将各种格式的文档转换为Markdown，尽可能保留原文档的所有内容
"""

import os
from typing import Optional


class DocumentConverter:
    """文档格式转换器"""

    def __init__(self, config: dict = None):
        self.config = config or {}

    def convert(self, input_path: str, output_path: str) -> bool:
        """将文档转换为Markdown格式"""
        if not os.path.exists(input_path):
            return False

        ext = os.path.splitext(input_path)[1].lower()

        try:
            if ext == '.docx':
                return self._convert_docx(input_path, output_path)
            elif ext in ['.xlsx', '.xls']:
                return self._convert_excel(input_path, output_path)
            elif ext == '.pdf':
                return self._convert_pdf(input_path, output_path)
            elif ext == '.md':
                return self._copy_md(input_path, output_path)
            else:
                return False
        except Exception as e:
            print(f"转换失败: {e}")
            return False

    def convert_batch(self, input_files: list, output_dir: str) -> list:
        """批量转换文档"""
        os.makedirs(output_dir, exist_ok=True)
        output_files = []

        for input_file in input_files:
            base_name = os.path.splitext(os.path.basename(input_file))[0]
            output_path = os.path.join(output_dir, f"{base_name}.md")
            if self.convert(input_file, output_path):
                output_files.append(output_path)

        return output_files

    def _convert_docx(self, input_path: str, output_path: str) -> bool:
        """转换Word文档为Markdown - 简单模式，保留原文档结构"""
        try:
            from docx import Document

            doc = Document(input_path)
            md_lines = []

            # 按顺序处理文档中的所有元素
            for element in doc.element.body:
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

                if tag == 'p':
                    # 处理段落
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break

                    if para:
                        text = para.text
                        if para.style and para.style.name:
                            style_name = para.style.name
                            # 保留标题级别
                            if 'Heading' in style_name:
                                level = style_name.replace('Heading', '').strip()
                                if level.isdigit():
                                    level = int(level)
                                    text = f"{'#' * level} {text}"
                        md_lines.append(text)
                        md_lines.append("")  # 段落间空行

                elif tag == 'tbl':
                    # 处理表格
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break

                    if table and table.rows:
                        # 添加表格前的空行
                        if md_lines and md_lines[-1].strip():
                            md_lines.append("")

                        # 处理表格的每一行
                        for row_idx, row in enumerate(table.rows):
                            cells = []
                            for cell in row.cells:
                                # 清理单元格内容，但保留基本结构
                                cell_text = cell.text.strip()
                                # 处理换行符
                                cell_text = cell_text.replace('\n', ' ')
                                cells.append(cell_text)

                            # 转换为Markdown表格行
                            md_lines.append('| ' + ' | '.join(cells) + ' |')

                            # 第一行后添加分隔线
                            if row_idx == 0:
                                separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                                md_lines.append(separator)

                        md_lines.append("")  # 表格后空行

            # 写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))

            return True
        except Exception as e:
            print(f"Word转换失败: {e}")
            return False

    def _convert_excel(self, input_path: str, output_path: str) -> bool:
        """转换Excel文档为Markdown - 简单模式"""
        try:
            import pandas as pd

            md_lines = []

            # 读取所有工作表
            excel_file = pd.ExcelFile(input_path)
            for sheet_name in excel_file.sheet_names:
                md_lines.append(f"# {sheet_name}\n")

                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                if not df.empty:
                    # 表头
                    headers = [str(col) for col in df.columns]
                    md_lines.append('| ' + ' | '.join(headers) + ' |')
                    md_lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')

                    # 数据行
                    for _, row in df.iterrows():
                        cells = [str(val) if pd.notna(val) else '' for val in row]
                        # 清理换行符
                        cells = [cell.replace('\n', ' ') for cell in cells]
                        md_lines.append('| ' + ' | '.join(cells) + ' |')

                    md_lines.append("")

            # 写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))

            return True
        except Exception as e:
            print(f"Excel转换失败: {e}")
            return False

    def _convert_pdf(self, input_path: str, output_path: str) -> bool:
        """转换PDF文档为Markdown"""
        try:
            import pdfplumber

            md_lines = []

            with pdfplumber.open(input_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    md_lines.append(f"# Page {page_num}\n")

                    # 提取文本
                    text = page.extract_text()
                    if text:
                        md_lines.append(text)
                        md_lines.append("")

                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            for row_idx, row in enumerate(table):
                                if row:
                                    cells = [str(cell) if cell else '' for cell in row]
                                    # 清理换行符
                                    cells = [cell.replace('\n', ' ') for cell in cells]
                                    md_lines.append('| ' + ' | '.join(cells) + ' |')

                                    if row_idx == 0:
                                        separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                                        md_lines.append(separator)

                            md_lines.append("")

            # 写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))

            return True
        except Exception as e:
            print(f"PDF转换失败: {e}")
            return False

    def _copy_md(self, input_path: str, output_path: str) -> bool:
        """复制Markdown文件"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)

            return True
        except Exception as e:
            print(f"MD文件复制失败: {e}")
            return False
