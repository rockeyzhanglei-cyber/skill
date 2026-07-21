#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档解析器
使用python-docx解析Word文档
"""

import os
import re
from typing import Dict, List
from .base import BaseParser, ParsedDocument, Table, Field


class WordParser(BaseParser):
    """Word文档解析器"""

    def __init__(self, config: Dict = None):
        super().__init__(config)
        self.name = "word_parser"

    def can_parse(self, file_path: str) -> bool:
        """判断是否可以解析该文件"""
        ext = self.get_file_extension(file_path)
        return ext in ['.docx', '.doc']

    def parse(self, file_path: str) -> ParsedDocument:
        """解析Word文件"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")

        doc = Document(file_path)
        parsed_doc = ParsedDocument(source_file=file_path)

        # 解析表格
        for idx, table in enumerate(doc.tables):
            parsed_table = self._parse_table(table, idx)
            if parsed_table and parsed_table.fields:
                parsed_doc.tables.append(parsed_table)

        return parsed_doc

    def _parse_table(self, table, idx: int) -> Table:
        """解析单个表格"""
        # 转换为二维数组
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        if not data:
            return None

        # 检测是否是数据标准表格
        headers = data[0] if data else []

        # 常见的表头模式
        field_indicators = ['字段', '数据元', '名称', '标识', '类型', '长度', '约束']
        header_match_count = sum(1 for h in headers if any(ind in h for ind in field_indicators))

        # 如果表头匹配的字段少于2个，可能不是数据标准表格
        if header_match_count < 2:
            return None

        # 创建表
        table_obj = Table(name=f"表 {idx + 1}", comment=f"表 {idx + 1}")

        # 解析数据行
        for row_data in data[1:]:
            if len(row_data) >= len(headers):
                field = self._parse_field_row(headers, row_data)
                if field:
                    table_obj.fields.append(field)

        return table_obj

    def _parse_field_row(self, headers: List[str], cells: List[str]) -> Field:
        """从表格行解析字段"""
        field_dict = {}

        for idx, header in enumerate(headers):
            if idx >= len(cells):
                break

            cell = cells[idx]
            header_lower = header.lower()

            # 识别数据元标识/英文名
            if '数据元标识' in header or '标识' in header or '英文名' in header or '字段名' in header:
                field_dict['name'] = cell.strip()

            # 识别数据元名称/中文名
            elif '数据元名称' in header or '名称' in header or '中文名' in header or '字段' == header:
                comment = cell.strip()
                if comment.startswith('*'):
                    comment = comment[1:]
                    field_dict['constraint'] = 'M'
                field_dict['comment'] = comment

            # 识别约束
            elif '约束' in header or '必填' in header:
                constraint_str = cell.strip()
                if constraint_str == 'M' or '必填' in constraint_str:
                    field_dict['constraint'] = 'M'
                elif constraint_str == 'C' or '条件' in constraint_str:
                    field_dict['constraint'] = 'C'
                else:
                    field_dict['constraint'] = 'O'

            # 识别数据类型
            elif '数据类型' in header or '类型' in header:
                field_dict['field_type'] = cell.strip()

            # 识别表示格式/长度
            elif '表示格式' in header or '格式' in header or '长度' in header:
                length_str = cell.strip()
                match = re.search(r'(\d+)', length_str)
                if match:
                    field_dict['length'] = int(match.group(1))

            # 识别说明
            elif '说明' in header or '描述' in header:
                field_dict['description'] = cell.strip()

            # 识别数据元标识符
            elif '数据元' in header and '标识符' in header:
                field_dict['data_element_id'] = cell.strip()

        # 必须有英文名
        if 'name' not in field_dict or not field_dict['name']:
            return None

        # 默认约束为O
        if 'constraint' not in field_dict:
            field_dict['constraint'] = 'O'

        return Field(
            name=field_dict.get('name', ''),
            comment=field_dict.get('comment', ''),
            field_type=field_dict.get('field_type', ''),
            length=field_dict.get('length', 0),
            constraint=field_dict.get('constraint', 'O'),
            description=field_dict.get('description', ''),
            data_element_id=field_dict.get('data_element_id', '')
        )
